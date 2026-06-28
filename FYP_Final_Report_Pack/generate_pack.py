from __future__ import annotations

import json
import re
import shutil
import subprocess
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from pypdf import PdfReader


ROOT = Path(r"C:\Users\jeysa\Desktop\Final Years")
PACK = ROOT / "FYP_Final_Report_Pack"
SOURCE = PACK / "01_source_materials"
EVIDENCE = PACK / "02_project_evidence"
DRAFT = PACK / "03_report_draft"
TEMPLATE = Path(r"C:\Users\jeysa\Downloads\Revised-FYP-Template-v20250710 (1).docx")
SAMPLE = Path(r"C:\Users\jeysa\Desktop\1.Megat_FinalReport.pdf")
OBJECTIVE = Path(
    r"C:\Users\jeysa\.codex\attachments\2265380a-1c7c-43f9-84d2-96b75540949a\goal-objective.md"
)


def run(cmd: list[str], cwd: Path = ROOT) -> str:
    try:
        return subprocess.check_output(cmd, cwd=cwd, text=True, stderr=subprocess.STDOUT).strip()
    except Exception as exc:
        return f"unavailable: {exc}"


def read(path: Path, limit: int | None = None) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text if limit is None else text[:limit]


def safe_json(path: Path):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def inventory() -> list[Path]:
    ignored = {
        ".git",
        "node_modules",
        "dist",
        "__pycache__",
        ".venv",
        "FYP_Final_Report_Pack",
    }
    files: list[Path] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        if any(part in ignored for part in rel.parts):
            continue
        files.append(rel)
    return sorted(files, key=lambda p: str(p).lower())


def extract_template() -> tuple[list[str], list[str]]:
    doc = Document(TEMPLATE)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    toc_like = [
        p
        for p in paragraphs
        if re.match(r"^(CHAPTER\s+\d+|[1-6](?:\.\d+){0,2}\s+|DECLARATION|ABSTRACT|TABLE OF CONTENTS)", p)
    ]
    return paragraphs[:80], toc_like[:80]


def extract_sample() -> tuple[int, list[str], list[str]]:
    reader = PdfReader(str(SAMPLE))
    headings: list[str] = []
    first_pages: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if i < 8:
            first_pages.append(text[:1800])
        for raw in text.splitlines():
            line = " ".join(raw.split())
            if not line or len(line) > 120:
                continue
            if re.match(r"^(CHAPTER\s+\d+|[1-6](?:\.\d+){1,2}\s+|ABSTRACT|ACKNOWLEDG|DECLARATION|REFERENCES|APPENDIX)", line, re.I):
                headings.append(line)
    return len(reader.pages), headings[:120], first_pages


def evidence_counts(files: list[Path]) -> dict[str, object]:
    suffix_counts = Counter(p.suffix.lower() or "[no extension]" for p in files)
    important = [
        p
        for p in files
        if p.suffix.lower() in {".md", ".json", ".jsonl", ".py", ".tsx", ".ts", ".mjs", ".ps1", ".pdf", ".docx"}
    ]
    return {
        "file_count": len(files),
        "top_extensions": suffix_counts.most_common(20),
        "important_files": [str(p) for p in important[:250]],
    }


def build_markdown(counts: dict[str, object], template_toc: list[str], sample_pages: int) -> str:
    generated = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    qa_report = read(ROOT / "QA_ACCURACY_API_AUDIT_PROJECT_TREE.md", 9000)
    cleanup = read(ROOT / "FYP_MASTER_BACKUP_CLEANUP_REPORT.md", 6000)
    frontend_readme = read(ROOT / "frontend" / "README.md", 6000)

    return f"""# FYP Final Report Draft Pack

Generated: {generated}

Project title placeholder: **AI Avatar for Academic Advising using a Fine-Tuned Large Language Model**

## How to Use This Pack

Use the Word draft as the editable report base. Use the evidence files when filling figures, tables, and citations. The sample A-grade report is used for structure only. Do not copy its prose.

## Source Materials

- MMU report guide/objective: `{OBJECTIVE}`
- MMU report template copy: `{SOURCE / TEMPLATE.name}`
- A-grade sample report copy: `{SOURCE / SAMPLE.name}`
- Sample report length: {sample_pages} pages
- Installed writing skill: `C:\\Users\\jeysa\\.codex\\skills\\stop-slop`
- Autoresearch reference: `https://github.com/karpathy/autoresearch`

## Formatting Rules Applied

- Times New Roman, 12 pt body text.
- Chapter and section headings use 12 pt bold.
- Page margins: top 20 mm, bottom 40 mm, left 40 mm, right 25 mm.
- Body paragraphs use 1.5 line spacing and justified alignment.
- References should use IEEE numeric style.

## Draft Abstract

Academic advising in programme-heavy faculties depends on accurate answers to course structures, prerequisites, assessment information, and progression rules. Students often ask these questions through informal channels, and staff must repeat the same explanations across cohorts. This project develops Hive, an AI academic advising kiosk that combines a fine-tuned language model, retrieval-augmented generation, deterministic course guards, and an animated Ebee avatar interface for Multimedia University Faculty of Artificial Intelligence and Engineering advising.

The system uses a React and Vite frontend, a FastAPI backend, FAISS-based retrieval indexes, structured course knowledge files, and a Qwen/Unsloth fine-tuned runtime in WSL. The frontend sends student questions to the backend, displays grounded answers, and presents an avatar-based kiosk experience. The backend routes questions through exact QA matching, course-structure guards, hybrid retrieval, reranking, and fallback generation. The knowledge base includes Intelligent Robotics and Applied AI programme structures, course catalogues, prerequisite rules, QA pairs, and generated indexes.

Evaluation focused on factual answer accuracy and kiosk readiness. The latest project evidence records a full QA-pair live audit of 1315/1315 passing rows, product RAG evaluation of 23/23 passed, raw backend evaluation of 18/18 passed, backend tests with 7 passed, and kiosk readiness with both frontend and backend ready. These results show that deterministic course routing and exact QA matching reduced wrong answers for high-risk academic facts. The project concludes that a guarded RAG design is more suitable for academic advising than ungrounded text generation, especially where programme rules must match official source documents.

## Chapter 1: Introduction

### 1.1 Overview

Hive is an AI academic advising kiosk for students who need fast answers about programme structures, prerequisites, course information, and academic progression. The project combines a conversational interface with a domain knowledge base and an avatar presentation layer. The target setting is the Faculty of Artificial Intelligence and Engineering, where students may need advising support outside office hours or before meeting an academic advisor.

### 1.2 Problem Statements

Students can receive wrong advice if a chatbot answers from general language-model memory instead of official course data. Course plans and prerequisite rules also contain small details that matter, such as term placement, course codes, and project sequencing. A generic chatbot may sound confident while mixing facts from unrelated programmes. The project therefore needs factual grounding, deterministic guards for known academic rules, and repeatable evaluation against stored source answers.

### 1.3 Objectives

1. Build a kiosk-style AI academic advisor that answers student questions through a web interface.
2. Create a structured knowledge base from programme, course, and prerequisite data.
3. Combine retrieval, exact QA matching, and deterministic guards to reduce wrong academic answers.
4. Provide an avatar-based interaction layer for a more approachable advising experience.
5. Verify the system through repeatable accuracy, build, and readiness checks.

### 1.4 Scope

The project focuses on academic advising questions for selected FAIE programme data, especially Intelligent Robotics and Applied AI artefacts found in the workspace. The system does not replace official academic advisors. It provides first-line guidance and should direct students to official faculty channels for cases that require approval, appeals, or personal academic records.

## Chapter 2: Literature Review

### 2.1 Conversational Academic Advising

Academic advising chatbots must handle factual questions, not only open conversation. A student may ask about prerequisites, credit hours, trimester placement, or project requirements. The system must answer from verified data because an attractive interface cannot compensate for wrong academic facts.

### 2.2 Retrieval-Augmented Generation

Retrieval-augmented generation supports domain-specific question answering by retrieving relevant records before the model writes an answer. In this project, the knowledge base includes JSONL files, course structures, source facts, and FAISS indexes. The design limits unsupported generation by routing high-risk questions to exact source data.

### 2.3 Fine-Tuned Language Models

Fine-tuning adapts a base language model to the style and task patterns of the project. The WSL runtime keeps the fine-tuned model output available while deterministic RAG paths handle official academic facts. This separation lets the model support conversation while the retrieval and guard layers protect facts.

### 2.4 Avatar-Based Interfaces

The frontend uses Ebee avatar assets to make the kiosk feel like an advising station rather than a plain text form. The current main UI uses generated video and image-based avatar states. The older live animated avatar pipeline remains in a standalone showcase for future development.

## Chapter 3: Details of the Design

### 3.1 System Architecture

The system contains three main parts: the React kiosk frontend, the FastAPI backend, and the WSL fine-tuned model runtime. The frontend displays the chat interface and avatar. The backend exposes health and chat endpoints, manages routing, reads knowledge base files, and retrieves relevant context. The WSL runtime serves the fine-tuned Qwen/Unsloth model.

### 3.2 Knowledge Base

The knowledge base contains programme structures, prerequisite graphs, course catalogues, QA pairs, source facts, and generated course-knowledge files. The latest evidence records 1315 Intelligent Robotics QA pairs, 1492 Hive course QA pairs, 94 programme-structure rows, and multiple FAISS indexes for global and detailed retrieval.

### 3.3 Answer Routing

The backend checks exact QA-pair questions before broader routing. This prevents generic rules from answering a question meant for a specific source row. Deterministic course guards handle high-risk facts. Hybrid retrieval and reranking supply context for broader questions. Fallback generation only handles questions outside the deterministic paths.

### 3.4 Frontend Design

The frontend is a Vite React app. It supports a commercial kiosk launch flow, a deterministic course-knowledge guard, avatar readiness checks, and RAG evaluation scripts. The main avatar component uses `AvatarExact` with generated media assets under `frontend\\public\\avatar\\exact`.

## Chapter 4: Presentation of Data

### 4.1 Accuracy Results

The project evidence records these results:

- Full QA-pair live audit: 1315/1315 passed.
- Product RAG evaluation: 23/23 passed.
- Raw backend evaluation: 18/18 passed.
- Frontend build: passed.
- Backend tests: 7 passed.
- Kiosk readiness: frontend ready and backend ready.

### 4.2 Project Artefacts

The project contains {counts["file_count"]} report-relevant files after excluding build and cache folders. Key artefacts include frontend source files, backend API and RAG modules, knowledge-base JSONL files, FAISS indexes, evaluation reports, avatar assets, and handoff documents.

### 4.3 Git and Workspace State

The Git repository root resolves to `C:\\Users\\jeysa`, while the FYP folder is currently untracked inside that wider repository. Treat this FYP pack as a source snapshot rather than a clean project-level commit history. The remote configured in the parent repository is `https://github.com/Jetsaw/Kommu_Ai_ChatBot.git`.

## Chapter 5: Discussion on Findings

### 5.1 Accuracy Improvements

The largest accuracy gain came from exact QA matching and duplicate question removal. Earlier generated QA questions reused the same wording for different source answers. The backend could then return the first matching row, even when another source row was intended. The fix made source-specific questions and routed exact matches before broad rules.

### 5.2 Guarded RAG Versus Ungrounded Generation

The project results support a guarded RAG design for academic advising. The model can handle language variation, but official course facts need deterministic routing. A wrong prerequisite answer can affect student planning, so the system should prefer source-backed answers over fluent guesses.

### 5.3 Runtime Separation

The Windows workspace contains frontend and backend source files, while the active fine-tuned runtime lives in WSL under `Ubuntu-24.04:/home/jet/fyp-unsloth`. This separation helped preserve the working model environment but makes documentation important. The report should include a deployment diagram and a table of runtime paths.

### 5.4 Limitations

The current draft still needs supervisor details, official student details, final screenshots, final references, and any faculty-required declaration text. The system should also be re-tested immediately before submission because backend availability affects kiosk readiness and live RAG scores.

## Chapter 6: Conclusions and Recommendations

Hive demonstrates that an academic advising chatbot can answer programme-specific questions when the system combines source-grounded retrieval, deterministic guards, and repeatable evaluation. The current evidence shows strong factual accuracy on stored QA pairs and readiness checks. Future work should expand the official course dataset, add a supervisor-approved escalation policy, improve analytics for unanswered questions, and package the WSL backend deployment steps for repeatable installation.

## IEEE Reference Placeholders

[1] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," 2020.

[2] E. J. Hu et al., "LoRA: Low-Rank Adaptation of Large Language Models," 2021.

[3] J. Johnson, M. Douze, and H. Jegou, "Billion-scale similarity search with GPUs," IEEE Transactions on Big Data, 2019.

[4] FastAPI documentation.

[5] React documentation.

[6] Vite documentation.

[7] Official MMU FAIE course and programme documents used in the project knowledge base.

## Template Structure Snapshot

{chr(10).join(f"- {line}" for line in template_toc[:40])}

## Evidence Extracts

### Frontend README

```text
{frontend_readme}
```

### QA Accuracy Report Extract

```text
{qa_report}
```

### Cleanup Report Extract

```text
{cleanup}
```
"""


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(4.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(12)
        styles[style_name].font.bold = True


def add_markdown_to_docx(md: str, path: Path) -> None:
    doc = Document()
    style_document(doc)
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            p.paragraph_format.line_spacing = 1.5
        elif line.startswith("```"):
            continue
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
    doc.save(path)


def main() -> None:
    for directory in [SOURCE, EVIDENCE, DRAFT]:
        directory.mkdir(parents=True, exist_ok=True)

    shutil.copy2(TEMPLATE, SOURCE / TEMPLATE.name)
    shutil.copy2(SAMPLE, SOURCE / SAMPLE.name)
    shutil.copy2(OBJECTIVE, SOURCE / "goal-objective.md")

    files = inventory()
    counts = evidence_counts(files)
    template_preview, template_toc = extract_template()
    sample_pages, sample_headings, sample_first_pages = extract_sample()

    git_context = {
        "git_root": run(["git", "rev-parse", "--show-toplevel"]),
        "status_scoped_to_fyp": run(["git", "status", "--short", "--", str(ROOT)]),
        "remote": run(["git", "remote", "-v"]),
        "log": run(["git", "log", "--oneline", "--decorate", "--max-count=20"]),
    }

    (EVIDENCE / "project_file_inventory.json").write_text(
        json.dumps(counts, indent=2), encoding="utf-8"
    )
    (EVIDENCE / "git_context.json").write_text(
        json.dumps(git_context, indent=2), encoding="utf-8"
    )
    (EVIDENCE / "template_structure.md").write_text(
        "# Template Structure\n\n"
        + "\n".join(f"- {line}" for line in template_toc)
        + "\n\n## First Template Paragraphs\n\n"
        + "\n".join(f"- {line}" for line in template_preview),
        encoding="utf-8",
    )
    (EVIDENCE / "sample_report_structure.md").write_text(
        f"# Sample Report Structure\n\nPages: {sample_pages}\n\n"
        + "\n".join(f"- {line}" for line in sample_headings)
        + "\n\n## First Pages Extract\n\n"
        + "\n\n---\n\n".join(sample_first_pages),
        encoding="utf-8",
    )
    (EVIDENCE / "academic_integrity_note.md").write_text(
        "# Academic Integrity Note\n\n"
        "I did not add or use Turnitin bypass tools. Use this pack to write from project evidence, cite sources, and remove generic AI-style filler. "
        "The installed stop-slop skill is for cleaner prose, not for evading academic review.\n",
        encoding="utf-8",
    )
    (EVIDENCE / "autoresearch_reference.md").write_text(
        "# Autoresearch Reference\n\n"
        "Repository: https://github.com/karpathy/autoresearch\n\n"
        "This repository is not a Codex skill because it has no `SKILL.md`. It is useful as a process reference: define an experiment file, run one bounded experiment, log the metric, keep only changes that improve the metric, and discard failed attempts. For this FYP report, use that pattern for evaluation discipline, not as a dependency.\n",
        encoding="utf-8",
    )

    md = build_markdown(counts, template_toc, sample_pages)
    (DRAFT / "FYP_FINAL_REPORT_DRAFT.md").write_text(md, encoding="utf-8")
    add_markdown_to_docx(md, DRAFT / "FYP_FINAL_REPORT_DRAFT.docx")

    summary = {
        "pack": str(PACK),
        "source_materials": [p.name for p in SOURCE.iterdir()],
        "evidence_files": [p.name for p in EVIDENCE.iterdir()],
        "draft_files": [p.name for p in DRAFT.iterdir()],
    }
    (PACK / "README.md").write_text(
        "# FYP Final Report Pack\n\n"
        + json.dumps(summary, indent=2)
        + "\n\nOpen `03_report_draft/FYP_FINAL_REPORT_DRAFT.docx` for the editable Word draft.\n",
        encoding="utf-8",
    )
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
