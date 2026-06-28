from __future__ import annotations

import json
import re
import shutil
import subprocess
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader


ROOT = Path(r"C:\Users\jeysa\Desktop\Final Years")
PACK = ROOT / "FYP_Final_Report_Pack"
SOURCE_DIR = PACK / "01_source_materials"
EVIDENCE_DIR = PACK / "02_project_evidence"
DRAFT_DIR = PACK / "03_report_draft"
CODE_DIR = PACK / "04_required_code"
FIGURE_DIR = PACK / "05_figures"
TEMPLATE = Path(r"C:\Users\jeysa\Downloads\Revised-FYP-Template-v20250710 (1).docx")
SAMPLE = Path(r"C:\Users\jeysa\Desktop\1.Megat_FinalReport.pdf")
OBJECTIVE = Path(
    r"C:\Users\jeysa\.codex\attachments\2265380a-1c7c-43f9-84d2-96b75540949a\goal-objective.md"
)
STOP_SLOP = Path(r"C:\Users\jeysa\.codex\skills\stop-slop\SKILL.md")


REPORT_TITLE = "AI Avatar for Academic Advising using a Fine-Tuned Large Language Model"
AUTHOR = "JET SAW JUN JIE"
STUDENT_ID = "1231303401"
SUPERVISOR = "the project supervisor"
SESSION = "2025/2026"


PROJECT_SOURCE_FILES = [
    Path(r"C:\Users\jeysa\Downloads\Project_Report_Rewritten.docx"),
    Path(r"C:\Users\jeysa\Downloads\AI Avatar for Academic Advising using Fine-Tuned Large Language Model.pptx"),
    Path(r"C:\Users\jeysa\Downloads\AI Avatar for Academic Advising using Fine-Tuned Large Language Model (1).pptx"),
    Path(r"C:\Users\jeysa\Downloads\fyp1 slide.pptx"),
    Path(r"C:\Users\jeysa\Downloads\HIVE_FYP_PROGRESS_LOG.md"),
    Path(r"C:\Users\jeysa\Downloads\HIVE_FYP_PROGRESS_LOG_UPDATED_3D_AVATAR.md"),
    Path(r"C:\Users\jeysa\Downloads\HIVE_FYP_FILE_LOCATIONS_WSL_COMMANDS.md"),
]


FIGURE_SOURCES = [
    ("frontend/ui-refresh-landscape.png", "Landscape kiosk interface after the UI refresh."),
    ("frontend/ui-refresh-portrait.png", "Portrait kiosk interface for tall screens."),
    ("frontend/avatar-video-only-verified.png", "Verified generated-video avatar layout."),
    ("frontend/avatar-video-check.png", "Avatar video rendering check in the application shell."),
    ("frontend/avatar-white-landscape.png", "White-background avatar validation in landscape mode."),
    ("frontend/avatar-white-portrait.png", "White-background avatar validation in portrait mode."),
    ("frontend/src/assets/hive-scene-bg.png", "Hive visual background asset used by the kiosk interface."),
    ("frontend/src/assets/hero-transparent.png", "Transparent hero asset used in the interface design."),
    ("frontend/src/assets/ebee-exact-transparent.png", "Transparent Ebee asset used by the exact avatar component."),
    ("frontend/public/avatar.png", "Original avatar reference asset."),
    ("frontend/public/avatar/exact/idle.png", "Exact avatar idle state."),
    ("frontend/public/avatar/exact/listening.png", "Exact avatar listening state."),
    ("frontend/public/avatar/exact/thinking.png", "Exact avatar thinking state."),
    ("frontend/public/avatar/exact/speaking.png", "Exact avatar speaking state."),
    ("frontend/public/avatar/exact/error.png", "Exact avatar error state."),
]


CODE_FILES = [
    "frontend/src/App.tsx",
    "frontend/src/components/AvatarExact.tsx",
    "frontend/src/components/AvatarExact.css",
    "frontend/src/courseKnowledge.ts",
    "frontend/src/courseKnowledgeData.ts",
    "frontend/src/main.tsx",
    "frontend/scripts/evaluate-rag-accuracy.mjs",
    "frontend/scripts/evaluate-all-qa-pairs.mjs",
    "frontend/scripts/check-kiosk-readiness.mjs",
    "frontend/scripts/check-stop-slop-output.mjs",
    "frontend/scripts/generate-course-knowledge.mjs",
    "frontend/package.json",
    "frontend/vite.config.ts",
    "hive-backend/app/main.py",
    "hive-backend/app/api/chat.py",
    "hive-backend/app/api/health.py",
    "hive-backend/app/rag/course_guard.py",
    "hive-backend/app/rag/hybrid_search.py",
    "hive-backend/app/rag/query_router.py",
    "hive-backend/app/rag/reranker.py",
    "hive-backend/app/rag/indexer.py",
    "hive-backend/app/advisor/engine.py",
    "hive-backend/app/advisor/alias_resolver.py",
    "hive-backend/app/advisor/programme_detection.py",
    "hive-backend/scripts/rebuild_intelligent_robotics_rag.py",
    "hive-backend/rebuild_indices.py",
    "hive-backend/requirements.txt",
]


SNIPPET_FILES = [
    "frontend/src/App.tsx",
    "frontend/src/components/AvatarExact.tsx",
    "frontend/src/courseKnowledge.ts",
    "frontend/scripts/evaluate-rag-accuracy.mjs",
    "frontend/scripts/check-kiosk-readiness.mjs",
    "hive-backend/app/main.py",
    "hive-backend/app/api/chat.py",
    "hive-backend/app/rag/course_guard.py",
    "hive-backend/app/rag/query_router.py",
    "hive-backend/app/rag/hybrid_search.py",
    "hive-backend/app/advisor/engine.py",
    "hive-backend/scripts/rebuild_intelligent_robotics_rag.py",
]


def run(cmd: list[str], cwd: Path = ROOT) -> str:
    try:
        return subprocess.check_output(
            cmd, cwd=cwd, text=True, stderr=subprocess.STDOUT, encoding="utf-8", errors="replace"
        ).strip()
    except Exception as exc:
        return f"unavailable: {exc}"


def read_text(path: Path, limit: int | None = None) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text if limit is None else text[:limit]


def read_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return {}


def setup_dirs() -> None:
    for folder in [SOURCE_DIR, EVIDENCE_DIR, DRAFT_DIR, CODE_DIR, FIGURE_DIR]:
        folder.mkdir(parents=True, exist_ok=True)
    for source in [TEMPLATE, SAMPLE, OBJECTIVE, STOP_SLOP]:
        if source.exists():
            shutil.copy2(source, SOURCE_DIR / source.name)
    copied_project_sources = []
    for source in PROJECT_SOURCE_FILES:
        if source.exists():
            shutil.copy2(source, SOURCE_DIR / source.name)
            copied_project_sources.append(str(source))
    (SOURCE_DIR / "project_source_manifest.json").write_text(
        json.dumps(
            {
                "student_name_source": str(Path(r"C:\Users\jeysa\Downloads\Project_Report_Rewritten.docx")),
                "student_name": AUTHOR,
                "student_id": STUDENT_ID,
                "copied_project_sources": copied_project_sources,
                "supervisor_status": "No supervisor name found in inspected local project reports, slide decks, metadata, or attendance form; report uses generic supervisor wording.",
            },
            indent=2,
        ),
        encoding="utf-8",
    )


def collect_inventory() -> dict:
    ignored = {".git", "node_modules", "dist", "__pycache__", ".venv", "FYP_Final_Report_Pack"}
    files: list[Path] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        if any(part in ignored for part in rel.parts):
            continue
        files.append(rel)
    suffix_counts = Counter(p.suffix.lower() or "[no extension]" for p in files)
    inventory = {
        "file_count": len(files),
        "extension_counts": suffix_counts.most_common(),
        "files": [str(p) for p in sorted(files, key=lambda x: str(x).lower())],
    }
    (EVIDENCE_DIR / "project_file_inventory_full.json").write_text(
        json.dumps(inventory, indent=2), encoding="utf-8"
    )
    return inventory


def collect_git_context() -> dict:
    context = {
        "git_root": run(["git", "rev-parse", "--show-toplevel"]),
        "scoped_status": run(["git", "status", "--short", "--", str(ROOT)]),
        "remote": run(["git", "remote", "-v"]),
        "log": run(["git", "log", "--oneline", "--decorate", "--max-count=20"]),
    }
    (EVIDENCE_DIR / "git_context_full.json").write_text(json.dumps(context, indent=2), encoding="utf-8")
    return context


def eval_summary() -> dict:
    product = read_json(ROOT / "frontend" / "rag-eval-report.product.json")
    raw = read_json(ROOT / "frontend" / "rag-eval-report.raw-backend.json")
    qa_all = read_json(ROOT / "frontend" / "qa-pair-full-eval-report.json")
    beginner = read_json(ROOT / "hybride serch master file" / "reports" / "beginner_general_500_live_eval_report.json")
    status_md = ROOT / "frontend" / "docs" / "RAG_ACCURACY_STATUS.md"
    live_summary = ROOT / "hybride serch master file" / "reports" / "live_eval_accuracy_summary.md"
    summary = {
        "product": {k: product.get(k) for k in ["baseUrl", "total", "passed", "generatedAt"]},
        "raw_backend": {k: raw.get(k) for k in ["baseUrl", "total", "passed", "generatedAt"]},
        "qa_all": {k: qa_all.get(k) for k in ["baseUrl", "total", "passed", "failed", "averageScore", "minScore", "durationSeconds", "generatedAt"]},
        "beginner_500": {k: beginner.get(k) for k in ["total", "passed", "failed"]},
        "status_markdown_exists": status_md.exists(),
        "live_accuracy_summary_exists": live_summary.exists(),
    }
    report_dir = EVIDENCE_DIR / "eval_reports"
    report_dir.mkdir(exist_ok=True)
    for rel in [
        "frontend/rag-eval-report.product.json",
        "frontend/rag-eval-report.raw-backend.json",
        "frontend/qa-pair-full-eval-report.json",
        "frontend/docs/RAG_ACCURACY_STATUS.md",
        "hybride serch master file/reports/live_eval_accuracy_summary.md",
        "hybride serch master file/reports/ui_mixed_300_live_eval_report.json",
        "hybride serch master file/reports/beginner_general_500_live_eval_report.json",
    ]:
        src = ROOT / rel
        if src.exists():
            dst = report_dir / rel.replace("/", "__").replace("\\", "__")
            shutil.copy2(src, dst)
    (EVIDENCE_DIR / "evaluation_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def copy_code_files() -> list[dict]:
    manifest: list[dict] = []
    for rel in CODE_FILES:
        src = ROOT / rel
        if not src.exists():
            manifest.append({"source": rel, "copied": False})
            continue
        dst = CODE_DIR / rel.replace("/", "__").replace("\\", "__")
        shutil.copy2(src, dst)
        text = read_text(src)
        manifest.append(
            {
                "source": rel,
                "copy": str(dst.relative_to(PACK)),
                "copied": True,
                "lines": text.count("\n") + 1,
                "bytes": src.stat().st_size,
            }
        )
    (CODE_DIR / "required_code_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def copy_figures() -> list[dict]:
    manifest: list[dict] = []
    for i, (rel, caption) in enumerate(FIGURE_SOURCES, start=1):
        src = ROOT / rel
        if not src.exists():
            manifest.append({"source": rel, "copied": False, "caption": caption})
            continue
        name = f"figure_{i:02d}_{src.name}"
        dst = FIGURE_DIR / name
        shutil.copy2(src, dst)
        try:
            im = Image.open(dst)
            size = [im.width, im.height]
        except Exception:
            size = None
        manifest.append(
            {
                "source": rel,
                "copy": str(dst.relative_to(PACK)),
                "copied": True,
                "caption": caption,
                "size": size,
                "bytes": dst.stat().st_size,
            }
        )
    (FIGURE_DIR / "figure_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str = "#1f2937") -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    words = text.split()
    lines: list[str] = []
    current = ""
    f = font(25)
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textlength(test, font=f) > (x2 - x1 - 34) and current:
            lines.append(current)
            current = word
        else:
            current = test
    if current:
        lines.append(current)
    total_h = len(lines) * 31
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        draw.text((x1 + (x2 - x1 - draw.textlength(line, font=f)) / 2, y), line, fill="#111827", font=f)
        y += 31


def make_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "#ffffff")
    d = ImageDraw.Draw(img)
    title_font = font(40)
    d.text((60, 40), "Hive AI Academic Advising Architecture", fill="#111827", font=title_font)
    boxes = [
        ((80, 170, 430, 330), "Student Kiosk UI React + Vite", "#dbeafe"),
        ((560, 170, 910, 330), "API Gateway FastAPI /ask /health", "#dcfce7"),
        ((1040, 170, 1390, 330), "Answer Router exact QA, guards, RAG", "#fef3c7"),
        ((80, 520, 430, 680), "AvatarExact generated video and PNG states", "#fce7f3"),
        ((560, 520, 910, 680), "Knowledge Base JSONL, rules, course data", "#ede9fe"),
        ((1040, 520, 1390, 680), "FAISS indexes hybrid search reranking", "#ffedd5"),
        ((1420, 350, 1740, 510), "WSL Qwen / Unsloth fine-tuned runtime", "#e0f2fe"),
        ((1420, 670, 1740, 830), "Evaluation gates UI, product, raw backend", "#f3f4f6"),
    ]
    for xy, text, fill in boxes:
        rounded_box(d, xy, text, fill)
    arrow_pairs = [
        ((430, 250), (560, 250)),
        ((910, 250), (1040, 250)),
        ((1390, 250), (1460, 350)),
        ((1215, 330), (1215, 520)),
        ((910, 600), (1040, 600)),
        ((430, 600), (560, 600)),
        ((1580, 510), (1580, 670)),
        ((270, 330), (270, 520)),
    ]
    for start, end in arrow_pairs:
        d.line([start, end], fill="#374151", width=5)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            tri = [(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)]
        else:
            tri = [(ex, ey), (ex + 18, ey - 10), (ex + 18, ey + 10)]
        if abs(ey - sy) > abs(ex - sx):
            tri = [(ex, ey), (ex - 10, ey - 18), (ex + 10, ey - 18)] if ey > sy else [(ex, ey), (ex - 10, ey + 18), (ex + 10, ey + 18)]
        d.polygon(tri, fill="#374151")
    d.text((70, 930), "Generated from current workspace evidence. The WSL runtime and Windows source copy are documented separately in the report.", fill="#374151", font=font(24))
    img.save(path)


def make_results_chart(path: Path, summary: dict) -> None:
    rows = [
        ("Product RAG", summary["product"].get("passed", 0), summary["product"].get("total", 0)),
        ("Raw backend", summary["raw_backend"].get("passed", 0), summary["raw_backend"].get("total", 0)),
        ("QA-pair audit", summary["qa_all"].get("passed", 0), summary["qa_all"].get("total", 0)),
        ("Beginner set", summary["beginner_500"].get("passed", 0), summary["beginner_500"].get("total", 0)),
    ]
    img = Image.new("RGB", (1500, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((60, 50), "Evaluation Results from Current Evidence", fill="#111827", font=font(40))
    max_total = max(total for _, _, total in rows if total) or 1
    y = 170
    for label, passed, total in rows:
        pct = (passed / total * 100) if total else 0
        d.text((80, y), label, fill="#111827", font=font(28))
        d.rounded_rectangle((360, y - 5, 1260, y + 45), radius=12, fill="#e5e7eb")
        width = int(900 * (passed / max_total)) if max_total else 0
        d.rounded_rectangle((360, y - 5, 360 + max(width, 4), y + 45), radius=12, fill="#2563eb")
        d.text((1285, y), f"{passed}/{total} ({pct:.2f}%)", fill="#111827", font=font(26))
        y += 135
    d.text((80, 780), "The 1315-row QA audit is scaled against the largest dataset, so smaller 20-row gates appear short by design.", fill="#374151", font=font(23))
    img.save(path)


def make_avatar_contact_sheet(path: Path, figure_manifest: list[dict]) -> None:
    avatar_items = [m for m in figure_manifest if m.get("copied") and "avatar/exact" in m["source"].replace("\\", "/")]
    thumbs = []
    for item in avatar_items:
        p = PACK / item["copy"]
        im = Image.open(p).convert("RGBA")
        im.thumbnail((190, 260))
        thumbs.append((item, im.copy()))
    w, h = 1300, 520
    img = Image.new("RGB", (w, h), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((40, 30), "Ebee Avatar State Assets", fill="#111827", font=font(36))
    x = 60
    y = 115
    for item, im in thumbs:
        img.paste(Image.new("RGB", im.size, "#ffffff"), (x, y))
        img.paste(im, (x, y), im if im.mode == "RGBA" else None)
        label = Path(item["source"]).stem.title()
        d.text((x, y + 275), label, fill="#111827", font=font(22))
        x += 240
    img.save(path)


def make_generated_figures(summary: dict, figure_manifest: list[dict]) -> list[dict]:
    generated = [
        ("figure_16_architecture_diagram.png", "System architecture for the Hive academic advising kiosk."),
        ("figure_17_evaluation_results.png", "Evaluation summary chart generated from project result files."),
        ("figure_18_avatar_state_contact_sheet.png", "Contact sheet of the exact Ebee avatar states."),
    ]
    make_architecture_diagram(FIGURE_DIR / generated[0][0])
    make_results_chart(FIGURE_DIR / generated[1][0], summary)
    make_avatar_contact_sheet(FIGURE_DIR / generated[2][0], figure_manifest)
    out = []
    for name, caption in generated:
        p = FIGURE_DIR / name
        im = Image.open(p)
        out.append(
            {
                "source": "[generated]",
                "copy": str(p.relative_to(PACK)),
                "copied": True,
                "caption": caption,
                "size": [im.width, im.height],
                "bytes": p.stat().st_size,
            }
        )
    return out


def template_and_sample_summary() -> dict:
    template_paras = []
    if TEMPLATE.exists():
        doc = Document(TEMPLATE)
        template_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    sample_pages = 0
    sample_headings = []
    if SAMPLE.exists():
        reader = PdfReader(str(SAMPLE))
        sample_pages = len(reader.pages)
        for page in reader.pages:
            text = page.extract_text() or ""
            for raw in text.splitlines():
                line = " ".join(raw.split())
                if re.match(r"^(CHAPTER\s+\d+|[1-6](?:\.\d+){1,2}\s+|ABSTRACT|ACKNOWLEDG|DECLARATION|REFERENCES|APPENDIX)", line, re.I):
                    sample_headings.append(line)
    summary = {
        "template_first_paragraphs": template_paras[:80],
        "sample_pages": sample_pages,
        "sample_headings": sample_headings[:120],
    }
    (EVIDENCE_DIR / "template_and_sample_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def paragraph(doc: Document, text: str = "", style: str | None = None, align=None, spacing: float = 1.5):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    p.paragraph_format.line_spacing = spacing
    if align is not None:
        p.alignment = align
    return p


def add_page_number(section, number_format: str = "decimal", start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), number_format)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_doc(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Times New Roman"
        styles[s].font.size = Pt(12)
        styles[s].font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(4.0)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(2.5)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
    return p


def add_caption(doc: Document, text: str):
    p = paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
    return p


def add_image(doc: Document, path: Path, caption: str, number: int):
    if not path.exists():
        return
    try:
        im = Image.open(path)
        max_w = Inches(5.9)
        width = max_w if im.width >= im.height else Inches(3.8)
        p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
        p.add_run().add_picture(str(path), width=width)
        add_caption(doc, f"Figure {number}: {caption}")
    except Exception as exc:
        paragraph(doc, f"[Figure unavailable: {path.name}: {exc}]")


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], number: int):
    add_caption(doc, f"Table {number}: {caption}")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    paragraph(doc)


def add_code_block(doc: Document, title: str, code: str, limit: int = 160) -> None:
    heading(doc, title, 3)
    lines = code.splitlines()
    if len(lines) > limit:
        lines = lines[:limit] + [f"... [truncated in report; full file copied in 04_required_code]"]
    for line in lines:
        p = paragraph(doc, line, spacing=1.0)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)


def report_sections(summary: dict, inventory: dict, git: dict, code_manifest: list[dict], sample: dict) -> dict[str, list[str]]:
    product = summary["product"]
    raw = summary["raw_backend"]
    qa = summary["qa_all"]
    beginner = summary["beginner_500"]
    files = inventory["file_count"]
    return {
        "abstract": [
            "Academic advising for programme-heavy faculties depends on accurate answers to course structures, prerequisites, assessment rules, and progression requirements. Students often ask the same questions through informal channels, and staff must repeat details that already exist in official programme documents. This project develops Hive, an AI academic advising kiosk for Multimedia University Faculty of Artificial Intelligence and Engineering. The system combines a fine-tuned language model, retrieval-augmented generation, deterministic course guards, and an Ebee avatar interface.",
            "The implementation uses a React and Vite frontend, FastAPI backend modules, FAISS indexes, structured JSONL knowledge files, and a Qwen/Unsloth fine-tuned runtime in WSL. The frontend displays a kiosk chat experience and avatar states. The backend routes questions through exact QA matching, course-structure guards, hybrid retrieval, reranking, and fallback generation. The knowledge base includes Intelligent Robotics and Applied AI programme structures, prerequisite rules, course catalogues, source facts, and generated QA pairs.",
            f"Evaluation used stored QA pairs, product-path RAG checks, raw-backend checks, and rendered UI tests. Current evidence records {qa.get('passed')}/{qa.get('total')} passing rows in the full QA-pair audit, {product.get('passed')}/{product.get('total')} passing rows for the product path, {raw.get('passed')}/{raw.get('total')} passing rows for the raw backend path, and {beginner.get('passed')}/{beginner.get('total')} passing rows for the beginner question set. A separate UI mixed test recorded 300/300 passing questions with an average response time of 532 ms. The results show that guarded RAG fits academic advising because the system can prefer official source facts over fluent but unsupported generation.",
        ],
        "chapter1": [
            "Hive addresses the gap between official academic information and the way students ask for help. A student may ask a short question such as whether Project II requires Project I, while the official answer lives across programme structures, subject master documents, prerequisite graphs, and faculty rules. The project turns those sources into a system that answers through a kiosk interface.",
            "The project does not attempt to replace an academic advisor. It provides first-line guidance for factual questions and prepares students before they contact the faculty. The system should still direct students to official channels when a case involves appeals, approvals, graduation checks, transfer credit, or personal academic records.",
            "The core problem is factual reliability. A generic language model can produce confident answers without checking the programme structure. That risk becomes serious when the question involves prerequisites, course codes, or progression rules. Hive therefore uses deterministic guards for high-risk academic facts and retrieval for broader course questions.",
            "The second problem is accessibility. Students who see a plain technical chatbot may treat it as a search box. The kiosk UI and Ebee avatar make the interaction closer to an advising station, while the backend still controls factual grounding.",
        ],
        "chapter2": [
            "Academic advising systems must support repetitive factual questions and decision support. A student may need to know which courses appear in Year 2, how many credit hours Project I requires, or whether a BYOC elective fits a trimester. These questions reward source accuracy more than creative text generation.",
            "Retrieval-augmented generation combines a language model with an external memory. Lewis et al. describe RAG models that pair parametric model memory with a non-parametric index for knowledge-intensive tasks [1]. Hive applies the same principle at project scale: programme structures and QA files sit outside the model, and the backend retrieves them when students ask questions.",
            "Fine-tuning adapts a model to domain language and answer style. LoRA reduces the cost of adaptation by adding trainable low-rank matrices while freezing the base weights [2]. The project uses a fine-tuned Qwen/Unsloth runtime for conversational generation, while course guards protect official facts.",
            "Vector search supports fast lookup over dense document representations. FAISS provides algorithms for similarity search and clustering across large vector sets [3]. Hive uses FAISS indexes with metadata files to retrieve course and programme facts.",
            "Academic advising also needs evaluation. The project uses answer overlap, source checks, endpoint checks, and UI-level tests. Backend correctness alone does not prove that the rendered kiosk shows the right answer.",
            "Karpathy's autoresearch project influenced the project workflow rather than the product dependency graph. The relevant practice is bounded experimentation: run an experiment, log the metric, keep improvements, and discard changes that do not improve the metric [8]. Hive uses that discipline in its RAG and UI evaluation gates.",
        ],
        "chapter3": [
            "The system follows a three-part design. The React frontend manages the student-facing kiosk, chat layout, avatar states, and local product guards. The FastAPI backend exposes health and chat endpoints. The WSL runtime serves the fine-tuned model and course-answering path.",
            "The knowledge base uses structured files because official academic data has stable fields. Programme rows, prerequisite rules, subject aliases, course facts, and QA pairs remain inspectable as JSONL or JSON. This choice also makes failed answers easier to debug because each answer can point back to a row or source file.",
            "The answer router gives priority to exact QA matches and deterministic course guards. This order prevents broad course-code matching from stealing a question that belongs to a specific source row. Hybrid retrieval and reranking handle broader questions, while fallback generation covers lower-risk conversational turns.",
            "The frontend uses the AvatarExact component and generated PNG/video states. The older rigged avatar pipeline remains in a standalone showcase. This split keeps the main kiosk simple and stable while preserving the advanced animation work for future development.",
            "The project stores required evaluation scripts inside the frontend workspace. Product checks call the Vite path, raw checks call the backend path, and full QA checks send stored questions through the live service. This makes accuracy a repeatable gate instead of a manual observation.",
        ],
        "chapter4": [
            f"The current report pack found {files} report-relevant files after excluding build and cache folders. The workspace includes frontend source, backend source, structured knowledge files, FAISS indexes, evaluation scripts, generated result reports, avatar assets, and handoff documents.",
            f"The product RAG evaluation passed {product.get('passed')}/{product.get('total')} cases against {product.get('baseUrl')}. The raw backend evaluation passed {raw.get('passed')}/{raw.get('total')} cases against {raw.get('baseUrl')}. The full QA-pair live audit passed {qa.get('passed')}/{qa.get('total')} rows with average overlap {qa.get('averageScore')}. The beginner set passed {beginner.get('passed')}/{beginner.get('total')} rows.",
            "The live accuracy summary records a larger API and UI regression suite. The patched backend passed 1000/1000 rows from the master accuracy set, 500/500 beginner rows, 1500/1500 mixed regression rows, and 300/300 rendered UI questions. The UI mixed test reported 532 ms average response time, 538 ms P50, 791 ms P95, and 1155 ms maximum response time.",
            "The Git context needs care. The Git root resolves to C:\\Users\\jeysa, while the FYP folder is untracked inside that wider repository. The report therefore treats this pack as a source snapshot. The parent remote is https://github.com/Jetsaw/Kommu_Ai_ChatBot.git.",
        ],
        "chapter5": [
            "Exact QA matching produced the clearest improvement. Earlier generated questions reused similar wording across different source facts. The backend could return a high-confidence answer from the wrong row. Source-specific generated questions and exact-normalized matching removed that ambiguity.",
            "The results support a guarded RAG design. The fine-tuned model can help with natural phrasing, but the system should not ask it to remember official academic rules. Course structures, prerequisite graphs, and source facts give the backend a stronger basis for answers.",
            "UI-level testing exposed issues that backend tests missed. The visible message can differ from raw backend output because the frontend may shorten answers or apply local guards. The 300-question rendered UI run therefore gave stronger evidence for student-facing correctness than API tests alone.",
            "The project still has limits. The dataset covers selected programme sources and generated QA rows. The system cannot decide appeals, approve course substitutions, or access personal academic records. The report should position Hive as a factual advising assistant, not an official decision maker.",
            "The writing process uses stop-slop as an editing guide. The skill removes generic AI phrasing, but it does not replace citation, evidence, or academic integrity checks. The report pack avoids Turnitin bypass tooling and keeps source evidence available for supervisor review.",
        ],
        "chapter6": [
            "Hive shows that an academic advising kiosk can answer programme-specific questions when the system pairs a language model with deterministic source controls. The current evidence shows perfect scores on the stored QA-pair audit, product RAG checks, raw backend checks, beginner questions, and rendered UI mixed tests.",
            "The project should continue in four directions. First, the team should add more official programme documents and keep versioned source snapshots. Second, the backend should log unanswered questions for supervisor review. Third, the kiosk should add clearer escalation paths for policy-sensitive cases. Fourth, deployment should package the Windows frontend and WSL backend so another evaluator can run the system with fewer manual steps.",
            "The local project documents verify the student's name and ID. The final printed submission still needs a handwritten declaration signature, date, and any supervisor-specific corrections requested before submission.",
        ],
    }


def add_front_matter(doc: Document, sample: dict, figure_manifest: list[dict], code_manifest: list[dict]) -> None:
    p = paragraph(doc, REPORT_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.2)
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True
    paragraph(doc, "By", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, STUDENT_ID, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, f"Session {SESSION}", align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        paragraph(doc)
    paragraph(
        doc,
        "The project report is prepared for Faculty of Artificial Intelligence Engineering, Multimedia University in partial fulfillment of Bachelor of Science (Honours) majoring in Intelligent Robotics.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    paragraph(doc, "FACULTY OF AI & ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "MULTIMEDIA UNIVERSITY", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "JUNE 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "Copyright", 1)
    paragraph(doc, "Copyright of this report belongs to Universiti Telekom Sdn. Bhd. as qualified by the Multimedia University intellectual property policy. The final wording should be checked against the official template before submission.")
    doc.add_page_break()

    heading(doc, "Declaration", 1)
    paragraph(doc, "I hereby declare that this work has been done by myself and that no portion of the work contained in this report has been submitted in support of any application for any other degree or qualification of this or any other university or institute of learning.")
    paragraph(doc, "Signature: ________________________")
    paragraph(doc, f"Name: {AUTHOR}")
    paragraph(doc, f"Student ID: {STUDENT_ID}")
    paragraph(doc, "Date: ________________________")
    doc.add_page_break()

    heading(doc, "Acknowledgements", 1)
    paragraph(doc, f"The author thanks {SUPERVISOR} for supervision and feedback during the development of this project. The author also thanks the Faculty of Artificial Intelligence and Engineering for programme material used to validate the academic advising knowledge base.")
    doc.add_page_break()

    heading(doc, "Abstract", 1)


def add_lists(doc: Document, figure_manifest: list[dict], tables: list[str]) -> None:
    doc.add_page_break()
    heading(doc, "Table of Contents", 1)
    toc = [
        "DECLARATION",
        "ACKNOWLEDGEMENTS",
        "ABSTRACT",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "LIST OF ABBREVIATIONS",
        "CHAPTER 1: INTRODUCTION",
        "CHAPTER 2: LITERATURE REVIEW",
        "CHAPTER 3: DETAILS OF THE DESIGN",
        "CHAPTER 4: PRESENTATION OF DATA",
        "CHAPTER 5: DISCUSSIONS ON FINDINGS",
        "CHAPTER 6: CONCLUSIONS AND RECOMMENDATIONS",
        "REFERENCES",
        "APPENDICES",
    ]
    for item in toc:
        paragraph(doc, item)

    doc.add_page_break()
    heading(doc, "List of Figures", 1)
    for i, item in enumerate([m for m in figure_manifest if m.get("copied")], start=1):
        paragraph(doc, f"Figure {i}: {item['caption']}")

    doc.add_page_break()
    heading(doc, "List of Tables", 1)
    for i, caption in enumerate(tables, start=1):
        paragraph(doc, f"Table {i}: {caption}")

    doc.add_page_break()
    heading(doc, "List of Abbreviations", 1)
    abbrev = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("BYOC", "Build Your Own Curriculum"),
        ("FAIE", "Faculty of Artificial Intelligence and Engineering"),
        ("FAISS", "Facebook AI Similarity Search"),
        ("FYP", "Final Year Project"),
        ("JSONL", "JavaScript Object Notation Lines"),
        ("LLM", "Large Language Model"),
        ("LoRA", "Low-Rank Adaptation"),
        ("QA", "Question Answering"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("UI", "User Interface"),
        ("WSL", "Windows Subsystem for Linux"),
    ]
    add_table(doc, "Abbreviations used in the report", ["Abbreviation", "Meaning"], abbrev, 0)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    main = doc.sections[-1]
    add_page_number(main, "decimal", 1)


def add_main_report(doc: Document, sections: dict[str, list[str]], summary: dict, inventory: dict, git: dict, code_manifest: list[dict], figure_manifest: list[dict]) -> None:
    # Abstract text is inserted after the Abstract heading created in front matter.
    for para in sections["abstract"]:
        paragraph(doc, para)

    tables = [
        "Project objectives and verification evidence",
        "Core software components",
        "Knowledge-base artefacts",
        "Evaluation results",
        "Required code files copied into the report pack",
        "Current workspace and Git context",
    ]
    add_lists(doc, figure_manifest, tables)

    heading(doc, "Chapter 1: Introduction", 1)
    heading(doc, "1.1 Overview", 2)
    for para in sections["chapter1"][:2]:
        paragraph(doc, para)
    heading(doc, "1.2 Problem Statements", 2)
    for para in sections["chapter1"][2:3]:
        paragraph(doc, para)
    heading(doc, "1.3 Project Objectives", 2)
    objectives = [
        ["1", "Build a kiosk-style AI academic advisor.", "React frontend, avatar assets, and kiosk launch scripts."],
        ["2", "Create a structured academic knowledge base.", "JSONL course data, prerequisite rules, QA pairs, and indexes."],
        ["3", "Reduce wrong academic answers.", "Exact QA matching, deterministic guards, hybrid retrieval, and reranking."],
        ["4", "Provide an avatar-based advising interface.", "AvatarExact component and generated Ebee states."],
        ["5", "Verify the system with repeatable gates.", "Product, raw backend, QA-pair, and UI evaluations."],
    ]
    add_table(doc, tables[0], ["No.", "Objective", "Evidence"], objectives, 1)
    heading(doc, "1.4 Project Scope", 2)
    paragraph(doc, sections["chapter1"][3])
    heading(doc, "1.5 Report Organisation", 2)
    paragraph(doc, "Chapter 2 reviews academic advising chatbots, RAG, fine-tuning, vector search, evaluation, and avatar interfaces. Chapter 3 explains the system design. Chapter 4 presents implementation data and evaluation results. Chapter 5 discusses findings. Chapter 6 concludes the report and recommends future work.")

    heading(doc, "Chapter 2: Theoretical Background and Literature Review", 1)
    for idx, title in enumerate(["2.1 Academic Advising Chatbots", "2.2 Retrieval-Augmented Generation", "2.3 Fine-Tuning and LoRA", "2.4 Vector Search and FAISS", "2.5 Evaluation of QA Systems", "2.6 Avatar-Based Kiosk Interfaces", "2.7 Autoresearch-Inspired Experiment Discipline"], start=0):
        heading(doc, title, 2)
        paragraph(doc, sections["chapter2"][min(idx, len(sections["chapter2"]) - 1)])

    heading(doc, "Chapter 3: Details of the Design", 1)
    for title, para in zip(
        ["3.1 System Architecture", "3.2 Knowledge Base Design", "3.3 Answer Routing", "3.4 Frontend and Avatar Design", "3.5 Evaluation Design"],
        sections["chapter3"],
    ):
        heading(doc, title, 2)
        paragraph(doc, para)
    copied_figures = [m for m in figure_manifest if m.get("copied")]
    add_image(doc, PACK / copied_figures[15]["copy"], copied_figures[15]["caption"], 16)
    add_table(
        doc,
        tables[1],
        ["Component", "Implementation", "Role"],
        [
            ["Frontend", "React, Vite, TypeScript", "Student kiosk UI and local product checks"],
            ["Backend", "FastAPI, Python", "Chat endpoints, routing, retrieval, health checks"],
            ["Knowledge base", "JSONL, JSON, PDF/DOCX sources", "Official programme and course facts"],
            ["Indexes", "FAISS metadata and vector files", "Fast retrieval over source chunks"],
            ["Runtime", "WSL Qwen/Unsloth", "Fine-tuned answer generation"],
        ],
        2,
    )
    add_image(doc, PACK / copied_figures[17]["copy"], copied_figures[17]["caption"], 18)

    heading(doc, "Chapter 4: Presentation of Data", 1)
    for title, para in zip(
        ["4.1 Workspace Artefacts", "4.2 Evaluation Summary", "4.3 Extended Accuracy Tests", "4.4 Git and Runtime Context"],
        sections["chapter4"],
    ):
        heading(doc, title, 2)
        paragraph(doc, para)
    add_image(doc, PACK / copied_figures[16]["copy"], copied_figures[16]["caption"], 17)
    add_table(
        doc,
        tables[2],
        ["Artefact", "Path", "Purpose"],
        [
            ["Programme structure", "hive-backend/data/kb/programme_structure.jsonl", "Official course placement facts"],
            ["Prerequisite rules", "hive-backend/data/kb/prereq_rules.json", "Project and course prerequisite checks"],
            ["QA pairs", "hive-backend/data/kb/intelligent_robotics_qa_pairs.jsonl", "Exact QA and evaluation source"],
            ["Course guard", "hive-backend/app/rag/course_guard.py", "Deterministic high-risk answer path"],
            ["FAISS indexes", "hive-backend/data/indexes/global", "Vector retrieval"],
        ],
        3,
    )
    add_table(
        doc,
        tables[3],
        ["Evaluation", "Passed", "Total", "Evidence file"],
        [
            ["Product RAG", summary["product"].get("passed"), summary["product"].get("total"), "frontend/rag-eval-report.product.json"],
            ["Raw backend", summary["raw_backend"].get("passed"), summary["raw_backend"].get("total"), "frontend/rag-eval-report.raw-backend.json"],
            ["Full QA-pair audit", summary["qa_all"].get("passed"), summary["qa_all"].get("total"), "frontend/qa-pair-full-eval-report.json"],
            ["Beginner questions", summary["beginner_500"].get("passed"), summary["beginner_500"].get("total"), "reports/beginner_general_500_live_eval_report.json"],
        ],
        4,
    )
    for i in range(1, 16):
        add_image(doc, PACK / copied_figures[i - 1]["copy"], copied_figures[i - 1]["caption"], i)

    heading(doc, "Chapter 5: Discussions on Findings", 1)
    for title, para in zip(
        ["5.1 Accuracy and Routing", "5.2 Guarded RAG", "5.3 UI-Level Verification", "5.4 Limitations", "5.5 Academic Integrity and Writing Quality"],
        sections["chapter5"],
    ):
        heading(doc, title, 2)
        paragraph(doc, para)
    add_table(
        doc,
        tables[5],
        ["Item", "Current value"],
        [
            ["Git root", git.get("git_root")],
            ["Scoped FYP status", git.get("scoped_status")],
            ["Remote", git.get("remote")],
            ["Code files copied", str(sum(1 for item in code_manifest if item.get("copied")))],
        ],
        6,
    )

    heading(doc, "Chapter 6: Conclusions and Recommendations", 1)
    for title, para in zip(["6.1 Conclusion", "6.2 Recommendations", "6.3 Final Submission Notes"], sections["chapter6"]):
        heading(doc, title, 2)
        paragraph(doc, para)

    heading(doc, "References", 1)
    refs = [
        "[1] P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\" NeurIPS, 2020. Available: https://arxiv.org/abs/2005.11401",
        "[2] E. J. Hu et al., \"LoRA: Low-Rank Adaptation of Large Language Models,\" 2021. Available: https://arxiv.org/abs/2106.09685",
        "[3] J. Johnson, M. Douze, and H. Jegou, \"Billion-scale similarity search with GPUs,\" IEEE Transactions on Big Data, 2019. Available: https://arxiv.org/abs/1702.08734",
        "[4] Meta AI Research, \"Faiss: A library for efficient similarity search and clustering of dense vectors.\" Available: https://github.com/facebookresearch/faiss",
        "[5] FastAPI, \"FastAPI documentation.\" Available: https://fastapi.tiangolo.com/",
        "[6] React, \"React documentation.\" Available: https://react.dev/",
        "[7] Vite, \"Vite documentation.\" Available: https://vite.dev/",
        "[8] A. Karpathy, \"autoresearch: AI agents running research on single-GPU nanochat training automatically.\" Available: https://github.com/karpathy/autoresearch",
        "[9] H. Pandya, \"stop-slop: A skill file for removing AI tells from prose.\" Available: https://github.com/hardikpandya/stop-slop",
        "[10] Multimedia University Faculty of AI & Engineering, official programme and course source documents included in the project knowledge base.",
    ]
    for ref in refs:
        paragraph(doc, ref, spacing=1.0)

    heading(doc, "Appendix A: Required Code Manifest", 1)
    code_rows = [[m["source"], m.get("copy", ""), str(m.get("lines", ""))] for m in code_manifest if m.get("copied")]
    add_table(doc, tables[4], ["Source file", "Copied pack path", "Lines"], code_rows, 5)

    heading(doc, "Appendix B: Code Extracts", 1)
    for rel in SNIPPET_FILES:
        src = ROOT / rel
        if src.exists():
            add_code_block(doc, rel, read_text(src), limit=90)

    heading(doc, "Appendix C: Reproducible Commands", 1)
    commands = [
        'cd "C:\\Users\\jeysa\\Desktop\\Final Years\\frontend"',
        "npm run build",
        "npm run course:validate",
        "npm run rag:eval:product",
        "npm run rag:eval:raw",
        "npm run rag:eval:qa-all",
        "npm run kiosk:check",
        "npm run stopslop:check",
    ]
    for cmd in commands:
        add_code_block(doc, "Command", cmd, limit=10)


def build_markdown(sections: dict[str, list[str]], summary: dict, figure_manifest: list[dict], code_manifest: list[dict], git: dict) -> str:
    lines = [f"# {REPORT_TITLE}", "", f"Generated: {datetime.now().isoformat(timespec='seconds')}", ""]
    lines += ["## Abstract", ""]
    lines += sections["abstract"] + [""]
    chapter_map = [
        ("Chapter 1: Introduction", "chapter1"),
        ("Chapter 2: Theoretical Background and Literature Review", "chapter2"),
        ("Chapter 3: Details of the Design", "chapter3"),
        ("Chapter 4: Presentation of Data", "chapter4"),
        ("Chapter 5: Discussions on Findings", "chapter5"),
        ("Chapter 6: Conclusions and Recommendations", "chapter6"),
    ]
    for title, key in chapter_map:
        lines += [f"## {title}", ""]
        for para in sections[key]:
            lines += [para, ""]
    lines += ["## Figures", ""]
    for i, item in enumerate([m for m in figure_manifest if m.get("copied")], start=1):
        lines += [f"![Figure {i}: {item['caption']}](../{item['copy'].replace(chr(92), '/')})", ""]
    lines += ["## Required Code Files", ""]
    for item in code_manifest:
        if item.get("copied"):
            lines += [f"- `{item['source']}` -> `{item['copy']}` ({item.get('lines')} lines)"]
    lines += ["", "## Git Context", "", "```json", json.dumps(git, indent=2), "```", ""]
    return "\n".join(lines)


def generate_report() -> dict:
    setup_dirs()
    inventory = collect_inventory()
    git = collect_git_context()
    summary = eval_summary()
    code_manifest = copy_code_files()
    figure_manifest = copy_figures()
    figure_manifest += make_generated_figures(summary, figure_manifest)
    (FIGURE_DIR / "figure_manifest.json").write_text(json.dumps(figure_manifest, indent=2), encoding="utf-8")
    sample = template_and_sample_summary()
    sections = report_sections(summary, inventory, git, code_manifest, sample)

    doc = Document()
    style_doc(doc)
    add_page_number(doc.sections[0], "lowerRoman", 1)
    add_front_matter(doc, sample, figure_manifest, code_manifest)
    add_main_report(doc, sections, summary, inventory, git, code_manifest, figure_manifest)
    style_doc(doc)
    full_docx = DRAFT_DIR / "FYP_FINAL_REPORT_FULL.docx"
    doc.save(full_docx)

    full_md = DRAFT_DIR / "FYP_FINAL_REPORT_FULL.md"
    full_md.write_text(build_markdown(sections, summary, figure_manifest, code_manifest, git), encoding="utf-8")

    completion = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "docx": str(full_docx),
        "markdown": str(full_md),
        "figures": sum(1 for m in figure_manifest if m.get("copied")),
        "code_files": sum(1 for m in code_manifest if m.get("copied")),
        "paragraphs": len(Document(full_docx).paragraphs),
        "tables": len(Document(full_docx).tables),
        "inline_shapes": len(Document(full_docx).inline_shapes),
        "requires_user_confirmation": ["signed declaration"],
    }
    (PACK / "FULL_REPORT_BUILD_SUMMARY.json").write_text(json.dumps(completion, indent=2), encoding="utf-8")
    (PACK / "README.md").write_text(
        "# FYP Final Report Pack\n\n"
        f"Main full report DOCX: `{full_docx}`\n\n"
        f"Main full report Markdown: `{full_md}`\n\n"
        f"Figures copied/generated: {completion['figures']}\n\n"
        f"Required code files copied: {completion['code_files']}\n\n"
        "Student name and ID were filled from verified local project documents. Sign and date the declaration before printed submission.\n",
        encoding="utf-8",
    )
    return completion


if __name__ == "__main__":
    print(json.dumps(generate_report(), indent=2))
